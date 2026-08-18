"""Document Parser for MousaviTax Knowledge pipeline.

Migrated from MKA-Core (ADR-003).

Supports:
  - PDF (text-based) via pypdf + pdfplumber
  - DOCX via python-docx
  - Plain text / Markdown

Output is structured for downstream Embedding + Vector DB + Citation.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional


@dataclass
class ParsedChunk:
    """A contiguous text segment with optional page/section metadata."""

    text: str
    page: Optional[int] = None
    section: Optional[str] = None
    chunk_index: int = 0
    char_start: int = 0
    char_end: int = 0


@dataclass
class ParsedDocument:
    """Full parse result for one source document."""

    source_id: str
    title: str
    source_type: str  # pdf | docx | markdown | text
    full_text: str
    chunks: list[ParsedChunk] = field(default_factory=list)
    page_count: Optional[int] = None
    metadata: dict = field(default_factory=dict)
    error: Optional[str] = None

    @property
    def success(self) -> bool:
        return self.error is None and bool(self.full_text.strip())


class DocumentParser:
    """
    Extract clean text from knowledge documents.

    Usage:
        parser = DocumentParser()
        doc = parser.parse_file("قانون مالیاتهای مستقیم.pdf")
        # or
        doc = parser.parse_bytes(data, filename="x.pdf", source_id="drive-xxx")
    """

    def __init__(
        self,
        chunk_size: int = 1200,
        chunk_overlap: int = 200,
        min_chunk_chars: int = 50,
    ) -> None:
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_chars = min_chunk_chars

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def parse_file(
        self,
        path: str | Path,
        source_id: Optional[str] = None,
        title: Optional[str] = None,
    ) -> ParsedDocument:
        path = Path(path)
        if not path.exists():
            return ParsedDocument(
                source_id=source_id or path.name,
                title=title or path.stem,
                source_type="unknown",
                full_text="",
                error=f"File not found: {path}",
            )
        data = path.read_bytes()
        return self.parse_bytes(
            data,
            filename=path.name,
            source_id=source_id or path.stem,
            title=title or path.stem,
        )

    def parse_bytes(
        self,
        data: bytes,
        filename: str,
        source_id: str,
        title: Optional[str] = None,
    ) -> ParsedDocument:
        title = title or Path(filename).stem
        ext = Path(filename).suffix.lower().lstrip(".")

        try:
            if ext == "pdf":
                return self._parse_pdf(data, source_id, title)
            if ext in ("docx", "doc"):
                return self._parse_docx(data, source_id, title)
            if ext in ("md", "markdown", "txt"):
                text = data.decode("utf-8", errors="replace")
                return self._build_result(source_id, title, ext or "text", text)
            return ParsedDocument(
                source_id=source_id,
                title=title,
                source_type=ext or "unknown",
                full_text="",
                error=f"Unsupported format: .{ext}",
            )
        except Exception as e:
            return ParsedDocument(
                source_id=source_id,
                title=title,
                source_type=ext or "unknown",
                full_text="",
                error=str(e),
            )

    # ------------------------------------------------------------------
    # Format handlers
    # ------------------------------------------------------------------

    def _parse_pdf(self, data: bytes, source_id: str, title: str) -> ParsedDocument:
        # Prefer pdfplumber for better layout/Persian handling; fallback to pypdf
        pages_text: list[str] = []
        page_count = 0

        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(data)) as pdf:
                page_count = len(pdf.pages)
                for i, page in enumerate(pdf.pages, start=1):
                    t = page.extract_text() or ""
                    t = self._normalize_text(t)
                    if t.strip():
                        pages_text.append(t)
        except Exception:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            page_count = len(reader.pages)
            for page in reader.pages:
                t = page.extract_text() or ""
                t = self._normalize_text(t)
                if t.strip():
                    pages_text.append(t)

        full_text = "\n\n".join(pages_text)
        result = self._build_result(
            source_id, title, "pdf", full_text, page_count=page_count
        )

        # Attach page numbers to chunks when possible
        if pages_text and result.chunks:
            chunks: list[ParsedChunk] = []
            idx = 0
            for page_no, page_text in enumerate(pages_text, start=1):
                for ch in self._chunk_text(page_text):
                    ch.page = page_no
                    ch.chunk_index = idx
                    chunks.append(ch)
                    idx += 1
            result.chunks = chunks
        return result

    def _parse_docx(self, data: bytes, source_id: str, title: str) -> ParsedDocument:
        from docx import Document

        doc = Document(io.BytesIO(data))
        paragraphs = [
            self._normalize_text(p.text) for p in doc.paragraphs if p.text.strip()
        ]
        full_text = "\n\n".join(paragraphs)
        return self._build_result(source_id, title, "docx", full_text)

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _build_result(
        self,
        source_id: str,
        title: str,
        source_type: str,
        full_text: str,
        page_count: Optional[int] = None,
    ) -> ParsedDocument:
        full_text = self._normalize_text(full_text)
        chunks = self._chunk_text(full_text)
        return ParsedDocument(
            source_id=source_id,
            title=title,
            source_type=source_type,
            full_text=full_text,
            chunks=chunks,
            page_count=page_count,
            metadata={"char_count": len(full_text), "chunk_count": len(chunks)},
        )

    def _normalize_text(self, text: str) -> str:
        if not text:
            return ""
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _chunk_text(self, text: str) -> list[ParsedChunk]:
        if not text or len(text) < self.min_chunk_chars:
            if text.strip():
                return [
                    ParsedChunk(
                        text=text.strip(),
                        chunk_index=0,
                        char_start=0,
                        char_end=len(text),
                    )
                ]
            return []

        chunks: list[ParsedChunk] = []
        start = 0
        idx = 0
        n = len(text)

        while start < n:
            end = min(start + self.chunk_size, n)
            if end < n:
                window = text[start:end]
                break_at = max(
                    window.rfind("\n\n"),
                    window.rfind(".\n"),
                    window.rfind("؟"),
                    window.rfind("."),
                )
                if break_at > self.chunk_size // 3:
                    end = start + break_at + 1

            piece = text[start:end].strip()
            if len(piece) >= self.min_chunk_chars:
                chunks.append(
                    ParsedChunk(
                        text=piece,
                        chunk_index=idx,
                        char_start=start,
                        char_end=end,
                    )
                )
                idx += 1

            if end >= n:
                break
            start = max(end - self.chunk_overlap, start + 1)

        return chunks
